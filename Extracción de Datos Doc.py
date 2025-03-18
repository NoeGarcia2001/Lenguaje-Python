import os
from bs4 import BeautifulSoup
from docx import Document

base = "C:/Users/garci/Downloads/triddefs_xml/defs"
doc_file_path = "C:/Users/garci/Downloads/datos_extraidos.docx"

# Crear un nuevo documento de Word
doc = Document()
doc.add_heading('Datos Extraídos de Archivos XML', level=1)

for root, _, files in os.walk(base):
    for file in files:
        if file.endswith(".xml"):
            ruta_xml = os.path.join(root, file)
            print(f"\n📂 Procesando: {ruta_xml}")

            try:
                with open(ruta_xml, "r", encoding="utf-8") as f:
                    xml_contenido = f.read()

                soup = BeautifulSoup(xml_contenido, "lxml-xml")

                # Buscar dentro de <Info> si existe
                info = soup.find("Info")

                if info:
                    tipo_archivo = info.find("FileType")
                    extension = info.find("Ext")

                # Obtener solo la primera etiqueta <Bytes>
                bytes_tag = soup.find("Bytes")

                # Extraer los datos
                tipo_archivo_text = tipo_archivo.text.strip() if tipo_archivo else 'No encontrado'
                extension_text = extension.text if extension else 'No encontrada'
                bytes_text = bytes_tag.text.strip() if bytes_tag else 'No encontrada'

                print("📌 Datos extraídos:")
                print(f"- Tipo de archivo: {tipo_archivo_text}")
                print(f"- Extensión: {extension_text}")
                print(f"- Cabecera (Bytes): {bytes_text}")
                print("-" * 50)

                # Agregar los datos al documento de Word
                doc.add_paragraph(f"📂 Procesando: {ruta_xml}")
                doc.add_paragraph(f"- Tipo de archivo: {tipo_archivo_text}")
                doc.add_paragraph(f"- Extensión: {extension_text}")
                doc.add_paragraph(f"- Cabecera (Bytes): {bytes_text}")
                doc.add_paragraph("-" * 50)

            except Exception as e:
                print(f"❌ Error al leer {ruta_xml}: {e}")

# Guardar el documento de Word
doc.save(doc_file_path)
print(f"✅ Datos extraídos guardados en '{doc_file_path}'")
